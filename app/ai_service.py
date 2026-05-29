"""
ai_service.py v2.1
Impulse AI — DocFix

- analyze_document  : détection des titres via n8n/Gemini
- estimate_pages    : estimation des numéros de page via n8n/Gemini
- Fallback heuristique si n8n indisponible
"""

import json, re, logging, asyncio
import httpx
from enum import Enum
from dataclasses import dataclass
from typing import Optional

from docx import Document
from docx.shared import Pt

from .config import settings

logger = logging.getLogger("docfix.ai")

CHARS_PER_PAGE = 2800  # fallback si IA indisponible


# ── Types ─────────────────────────────────────────────────────────────────────
class AIProvider(str, Enum):
    N8N       = "n8n_gemini"
    HEURISTIC = "heuristic"


class AIErrorType(str, Enum):
    QUOTA_EXCEEDED = "quota_exceeded"
    INVALID_KEY    = "invalid_key"
    UNAVAILABLE    = "unavailable"
    TIMEOUT        = "timeout"
    NONE           = "none"


@dataclass
class AIResult:
    data:         dict
    provider:     AIProvider
    error_type:   AIErrorType
    error_detail: Optional[str] = None


# ── Helpers ───────────────────────────────────────────────────────────────────
def _extract_paragraphs(doc: Document) -> list:
    return [
        {"idx": i, "text": p.text.strip()}
        for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ][:50]


def _detect_error_type(msg: str) -> AIErrorType:
    m = msg.lower()
    if "429" in m or "quota" in m or "rate" in m: return AIErrorType.QUOTA_EXCEEDED
    if "401" in m or "403" in m or "key"  in m:   return AIErrorType.INVALID_KEY
    if "timeout" in m:                             return AIErrorType.TIMEOUT
    return AIErrorType.UNAVAILABLE


# ── Appel générique n8n ───────────────────────────────────────────────────────
async def _call_n8n(webhook_url: str, prompt: str) -> dict:
    headers = {"Content-Type": "application/json"}
    if settings.N8N_WEBHOOK_SECRET:
        headers["X-Webhook-Secret"] = settings.N8N_WEBHOOK_SECRET

    async with httpx.AsyncClient(timeout=40.0) as client:
        response = await client.post(webhook_url, headers=headers, json={"prompt": prompt})
        response.raise_for_status()
        data = response.json()
        if isinstance(data, list):
            data = data[0]
        return data


# ── 1. Analyse des titres ─────────────────────────────────────────────────────
def _build_analysis_prompt(paragraphs: list) -> str:
    return f"""Analyse ce document Word et identifie les titres et sous-titres.

Paragraphes (index: texte) :
{chr(10).join(f"[{p['idx']}] {p['text'][:120]}" for p in paragraphs)}

Réponds UNIQUEMENT en JSON :
{{
  "doc_type": "rapport|lettre|contrat|article|autre",
  "likely_titles": [
    {{"idx": 0, "level": 1, "reason": "court, sans ponctuation"}},
    {{"idx": 5, "level": 2, "reason": "numéroté, sous-section"}}
  ],
  "language": "fr|en|autre"
}}

Critères titre : <80 caractères, pas de point final, logique de section."""


def _heuristic_analysis(paragraphs: list) -> dict:
    titles = []
    for p in paragraphs:
        text = p["text"]
        if not text or len(text) > 100: continue
        no_dot = not text.endswith(".")
        short  = len(text) < 80
        num    = bool(__import__('re').match(r"^\d+[\.\)]\s+\w", text))
        caps   = text == text.upper() and len(text) > 3
        roman  = bool(__import__('re').match(r"^[IVX]+[\.\)]\s+\w", text))
        if short and no_dot and (num or caps or roman):
            titles.append({"idx": p["idx"], "level": 1 if caps else 2, "reason": "heuristique"})
    return {"doc_type": "unknown", "likely_titles": titles, "language": "fr"}


async def analyze_document(doc: Document, job_id: str = "") -> AIResult:
    paragraphs = _extract_paragraphs(doc)
    prompt     = _build_analysis_prompt(paragraphs)
    prefix     = f"[{job_id}] "

    if settings.N8N_GEMINI_WEBHOOK_URL:
        try:
            data  = await _call_n8n(settings.N8N_GEMINI_WEBHOOK_URL, prompt)
            count = len(data.get("likely_titles", []))
            logger.info(f"{prefix}✅ n8n/Gemini titres OK — {count} titres")
            return AIResult(data=data, provider=AIProvider.N8N, error_type=AIErrorType.NONE)
        except Exception as e:
            etype = _detect_error_type(str(e))
            logger.warning(f"{prefix}⚠️ n8n titres indisponible ({etype}) — fallback")
    else:
        logger.info(f"{prefix}ℹ️ N8N_GEMINI_WEBHOOK_URL non configurée")
        etype = AIErrorType.UNAVAILABLE

    data = _heuristic_analysis(paragraphs)
    logger.info(f"{prefix}🔧 Heuristique — {len(data.get('likely_titles',[]))} titres")
    return AIResult(data=data, provider=AIProvider.HEURISTIC, error_type=etype)


# ── 2. Estimation des numéros de page via IA ──────────────────────────────────
def _build_pages_prompt(doc: Document, heading_map: dict, start_page: int) -> str:
    """
    Construit le prompt pour estimer les pages.
    On envoie à Gemini :
    - La liste des titres avec leur position (index)
    - Le nombre de caractères entre chaque titre (volume de contenu)
    - Les paramètres de mise en page
    """
    # Collecter les titres et le volume entre eux
    entries = []
    prev_idx   = 0
    char_count = 0

    for i, para in enumerate(doc.paragraphs):
        # Accumuler les caractères
        char_count += len(para.text) + 1

        lvl = heading_map.get(i)
        if lvl:
            # Détecter saut de page explicite avant ce titre
            has_pb = any(
                br.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page"
                for br in para._element.iter(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
                )
            )
            entries.append({
                "idx":        i,
                "level":      min(lvl, 3),
                "text":       para.text.strip()[:80],
                "chars_before": char_count,
                "page_break": has_pb,
            })
            char_count = 0

    prompt = f"""Tu dois estimer les numéros de page de chaque titre dans un document Word.

Paramètres du document :
- Page de départ : {start_page}
- Police : Calibri 11pt
- Marges : 2.5cm sur 4 côtés (format A4)
- Caractères moyens par page : ~2800

Titres du document (dans l'ordre) :
{json.dumps(entries, ensure_ascii=False, indent=2)}

Pour chaque titre, estime le numéro de page en tenant compte :
- Du volume de texte (chars_before) avant ce titre
- Des sauts de page explicites (page_break: true)
- De la page de départ ({start_page})
- Les titres H1 (level 1) commencent souvent sur une nouvelle page

Réponds UNIQUEMENT en JSON :
{{
  "page_numbers": {{
    "IDX_DU_TITRE": NUMERO_DE_PAGE,
    ...
  }},
  "total_pages_estimated": NOMBRE_TOTAL
}}

Exemple : {{"page_numbers": {{"0": 1, "15": 2, "32": 4}}, "total_pages_estimated": 8}}"""

    return prompt


async def estimate_page_numbers(
    doc: Document,
    heading_map: dict,
    start_page: int = 1,
    job_id: str = "",
) -> dict:
    """
    Estime les numéros de page de chaque titre.
    Utilise Gemini via n8n si disponible, sinon fallback mathématique.
    Retourne {para_index: page_number}.
    """
    prefix = f"[{job_id}] "

    # ── Tentative IA ──────────────────────────────────────────────
    if settings.N8N_GEMINI_WEBHOOK_URL:
        try:
            prompt = _build_pages_prompt(doc, heading_map, start_page)
            data   = await _call_n8n(settings.N8N_GEMINI_WEBHOOK_URL, prompt)

            raw_pages = data.get("page_numbers", {})
            # Convertir les clés string en int
            page_map = {int(k): int(v) for k, v in raw_pages.items()}

            logger.info(
                f"{prefix}✅ n8n/Gemini pages OK — "
                f"{len(page_map)} pages estimées, "
                f"total ~{data.get('total_pages_estimated', '?')} pages"
            )
            return page_map

        except Exception as e:
            logger.warning(f"{prefix}⚠️ n8n pages indisponible: {e} — fallback mathématique")

    # ── Fallback mathématique ─────────────────────────────────────
    logger.info(f"{prefix}🔧 Estimation mathématique des pages")
    return _math_estimate_pages(doc, heading_map, start_page)


def _math_estimate_pages(doc: Document, heading_map: dict, start_page: int) -> dict:
    """Estimation mathématique locale (fallback)."""
    from docx.oxml.ns import qn as _qn
    page_map     = {}
    char_count   = 0
    current_page = start_page

    for i, para in enumerate(doc.paragraphs):
        has_pb = any(
            br.get(_qn("w:type")) == "page"
            for br in para._element.iter(_qn("w:br"))
        )
        if has_pb:
            current_page += 1
            char_count    = 0

        if i in heading_map:
            page_map[i] = current_page

        char_count += len(para.text) + 1
        if char_count >= CHARS_PER_PAGE:
            current_page += char_count // CHARS_PER_PAGE
            char_count    = char_count % CHARS_PER_PAGE

        if heading_map.get(i) == 1 and char_count > CHARS_PER_PAGE * 0.85:
            current_page += 1
            char_count    = 0

    return page_map


# ── Status checks ─────────────────────────────────────────────────────────────
async def check_gemini_status() -> dict:
    if not settings.N8N_GEMINI_WEBHOOK_URL:
        return {"status": "not_configured", "provider": "n8n_gemini"}
    try:
        await _call_n8n(
            settings.N8N_GEMINI_WEBHOOK_URL,
            '[{"idx":0,"text":"Test"}]\nRéponds: {"doc_type":"test","likely_titles":[],"language":"fr"}'
        )
        return {"status": "ok", "provider": "n8n_gemini",
                "webhook": settings.N8N_GEMINI_WEBHOOK_URL}
    except Exception as e:
        return {"status": _detect_error_type(str(e)).value,
                "provider": "n8n_gemini", "detail": str(e)[:200]}


async def check_fallback_status() -> dict:
    return {"status": "disabled", "provider": "none"}
