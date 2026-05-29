"""
processor.py v2.0
Impulse AI — DocFix

Nouveautés :
- TOC avec numéros de page estimés
- start_page pris en compte dans l'estimation
"""

import asyncio, logging, subprocess, time, re, os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

from .config import settings
from .jobs import job_store
from .ai_service import analyze_document, AIProvider

logger = logging.getLogger("docfix.processor")

COLOR_H1     = RGBColor(0x1F, 0x35, 0x64)
COLOR_H2     = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H3     = RGBColor(0x40, 0x40, 0x40)
COLOR_BODY   = RGBColor(0x26, 0x26, 0x26)
COLOR_FOOTER = RGBColor(0x80, 0x80, 0x80)

HEX_H1       = "1F3564"
HEX_H2       = "2E74B5"
HEX_H3       = "404040"
HEX_BODY     = "262626"
HEX_FOOTER   = "808080"
HEX_TBL_HEAD = "1F3564"
HEX_TBL_ALT  = "EEF3FA"

FONT_HEAD = "Calibri"
FONT_BODY = "Calibri"

# Caractères moyens par page A4 (Calibri 11pt, marges 2.5cm)
CHARS_PER_PAGE = 2800


async def process_document(
    job_id:      str,
    input_path:  Path,
    start_page:  int  = 1,
    include_toc: bool = True,
):
    job = job_store.get(job_id)
    if not job:
        return

    output_docx = settings.TEMP_DIR / f"{job_id}_output.docx"
    output_pdf  = settings.TEMP_DIR / f"{job_id}_output.pdf"

    try:
        job.status = "processing"
        start = time.time()
        _step_done(job, "upload", 0)

        _step_start(job, "parse", "Analyse du document...")
        doc = Document(str(input_path))
        await asyncio.sleep(0)
        _step_done(job, "parse", 8)

        _step_start(job, "ai", "Analyse IA via n8n...")
        ai_result = await analyze_document(doc, job_id)
        _step_done(job, "ai", 18)

        _step_start(job, "clean", "Nettoyage espaces et pages blanches...")
        await asyncio.to_thread(_clean_whitespace, doc)
        await asyncio.to_thread(_remove_blank_pages, doc)
        _step_done(job, "clean", 26)

        _step_start(job, "fonts", "Harmonisation des polices...")
        await asyncio.to_thread(_harmonize_fonts, doc)
        _step_done(job, "fonts", 34)

        _step_start(job, "headings", "Detection des titres...")
        heading_map = await asyncio.to_thread(_detect_headings, doc, ai_result.data)
        n_heads = len(heading_map)
        logger.info(f"[{job_id}] {n_heads} titres detectes")
        _step_done(job, "headings", 44)

        _step_start(job, "styles", "Application styles, couleurs et TOC...")
        toc_ok = await asyncio.to_thread(
            _apply_all_formatting, doc, heading_map, include_toc, start_page
        )
        _step_done(job, "styles", 54)

        job.set_step_done("toc")
        job.recalculate_progress()
        job.progress = max(job.progress, 63)

        _step_start(job, "images", "Redimensionnement images...")
        n_img = await asyncio.to_thread(_fix_images, doc)
        _step_done(job, "images", 69)

        _step_start(job, "pagination", "Tableaux et pagination...")
        n_tables = await asyncio.to_thread(_format_tables, doc)
        await asyncio.to_thread(_add_pagination, doc, start_page)
        _step_done(job, "pagination", 79)

        _step_start(job, "export_docx", "Sauvegarde DOCX...")
        doc.save(str(output_docx))
        _step_done(job, "export_docx", 88)

        _step_start(job, "export_pdf", "Conversion PDF...")
        await asyncio.to_thread(_to_pdf, output_docx, output_pdf)
        _step_done(job, "export_pdf", 100)

        duration = round(time.time() - start, 1)
        job.stats = {
            "pagesCount":       _count_pages(doc),
            "headingsDetected": n_heads,
            "imagesFixed":      n_img,
            "tablesFormatted":  n_tables,
            "fontsHarmonized":  1,
            "durationSeconds":  duration,
            "aiProvider":       ai_result.provider.value,
            "tocInserted":      toc_ok,
            "startPage":        start_page,
        }
        job.status       = "done"
        job.progress     = 100
        job.current_step = "Traitement termine"
        logger.info(f"Job {job_id} OK en {duration}s — heads={n_heads} toc={toc_ok}")

    except Exception as exc:
        logger.error(f"Erreur job {job_id}: {exc}", exc_info=True)
        job.status       = "error"
        job.error        = str(exc)
        job.current_step = "Erreur de traitement"
        for s in job.steps:
            if s.status == "running":
                s.status = "error"
    finally:
        input_path.unlink(missing_ok=True)


def _apply_all_formatting(doc, heading_map, include_toc, start_page):
    _setup_styles(doc)
    _apply_heading_styles_with_map(doc, heading_map)
    if include_toc:
        return _add_toc_manual(doc, heading_map, start_page)
    return False


def _step_start(job, sid, msg):
    job.set_step_running(sid)
    job.current_step = msg
    logger.info(f"[{job.job_id}] {msg}")

def _step_done(job, sid, pct):
    job.set_step_done(sid)
    job.recalculate_progress()
    if pct > job.progress:
        job.progress = pct


# ── Nettoyage ─────────────────────────────────────────────────────────────────
def _clean_whitespace(doc):
    prev_empty = False
    to_del = []
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = re.sub(r' {2,}', ' ', run.text)
            run.text = re.sub(r'\t+', ' ', run.text)
        if not para.text.strip():
            if prev_empty:
                to_del.append(para)
            prev_empty = True
        else:
            prev_empty = False
    for p in to_del:
        p._element.getparent().remove(p._element)


def _remove_blank_pages(doc):
    page_break_paras = []
    for para in doc.paragraphs:
        if para.text.strip():
            continue
        for br in para._element.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                page_break_paras.append(para)
                break

    paras_list = list(doc.paragraphs)
    to_del = []
    for i, para in enumerate(paras_list):
        if para not in page_break_paras:
            continue
        if i > 0 and paras_list[i-1] in page_break_paras:
            to_del.append(para)

    for para in to_del:
        para._element.getparent().remove(para._element)

    if to_del:
        logger.info(f"Pages blanches supprimees: {len(to_del)}")

    # Supprimer paragraphes vides en fin de document
    paras = list(doc.paragraphs)
    to_del2 = []
    for para in reversed(paras):
        if not para.text.strip():
            to_del2.append(para)
        else:
            break
    for para in (to_del2[:-1] if len(to_del2) > 1 else []):
        para._element.getparent().remove(para._element)


# ── Polices ───────────────────────────────────────────────────────────────────
def _harmonize_fonts(doc):
    for para in doc.paragraphs:
        sn = para.style.name if para.style else ""
        if any(sn.startswith(x) for x in ["Heading", "Title", "TOC"]):
            continue
        for run in para.runs:
            if run.font.size and run.font.size > Pt(14):
                continue
            run.font.name = FONT_BODY
            if not run.font.size or not (Pt(8) <= run.font.size <= Pt(14)):
                run.font.size = Pt(11)
            _set_run_color(run, HEX_BODY)


# ── Détection titres ──────────────────────────────────────────────────────────
def _detect_headings(doc, ai_hints) -> dict:
    ai_map = {t["idx"]: t.get("level", 1) for t in ai_hints.get("likely_titles", [])}
    result = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) > 150:
            continue
        sn = para.style.name if para.style else ""
        if sn.startswith("Heading"):
            try:    lvl = int(sn.split()[-1])
            except: lvl = 1
            result[i] = lvl
            continue
        level = ai_map.get(i)
        if level is None:
            short  = len(text) < 80
            no_dot = not text.endswith(".")
            big_f  = any(r.font.size and r.font.size >= Pt(14) for r in para.runs if r.font.size)
            bold   = bool(para.runs) and all(r.bold for r in para.runs if r.text.strip())
            caps   = text == text.upper() and len(text) > 3
            num    = bool(re.match(r"^\d+[\.\)]\s+\w", text))
            roman  = bool(re.match(r"^[IVX]+[\.\)]\s+\w", text))
            if   big_f and short and no_dot:  level = 1
            elif bold  and short and no_dot:  level = 2
            elif (num  or roman) and short:   level = 2
            elif caps  and short and no_dot:  level = 1
        if level:
            result[i] = level
    logger.info(f"_detect_headings: {len(result)} titres")
    return result


# ── Setup styles ──────────────────────────────────────────────────────────────
def _setup_styles(doc):
    cfgs = {
        "Heading 1": dict(size=Pt(18), color=COLOR_H1, bold=True,  italic=False,
                          caps=True,  before=Pt(24), after=Pt(8),
                          align=WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 2": dict(size=Pt(14), color=COLOR_H2, bold=True,  italic=False,
                          caps=False, before=Pt(16), after=Pt(6),
                          align=WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": dict(size=Pt(12), color=COLOR_H3, bold=True,  italic=True,
                          caps=False, before=Pt(12), after=Pt(4),
                          align=WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, c in cfgs.items():
        try:
            try:    st = doc.styles[name]
            except: st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.font.name      = FONT_HEAD
            st.font.size      = c["size"]
            st.font.color.rgb = c["color"]
            st.font.bold      = c["bold"]
            st.font.italic    = c["italic"]
            if c["caps"]: st.font.all_caps = True
            pf = st.paragraph_format
            pf.space_before      = c["before"]
            pf.space_after       = c["after"]
            pf.alignment         = c["align"]
            pf.keep_with_next    = True
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except Exception as e:
            logger.warning(f"Style {name}: {e}")
    try:
        n = doc.styles["Normal"]
        n.font.name      = FONT_BODY
        n.font.size      = Pt(11)
        n.font.color.rgb = COLOR_BODY
        n.paragraph_format.space_after       = Pt(6)
        n.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
        n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    except: pass


# ── Appliquer styles ──────────────────────────────────────────────────────────
def _apply_heading_styles_with_map(doc, heading_map):
    hex_map   = {1: HEX_H1,  2: HEX_H2,  3: HEX_H3}
    size_map  = {1: Pt(18),  2: Pt(14),  3: Pt(12)}
    align_map = {1: WD_ALIGN_PARAGRAPH.CENTER,
                 2: WD_ALIGN_PARAGRAPH.LEFT,
                 3: WD_ALIGN_PARAGRAPH.LEFT}
    for i, para in enumerate(doc.paragraphs):
        lvl = heading_map.get(i)
        if not lvl: continue
        lvl = min(lvl, 3)
        try: para.style = doc.styles[f"Heading {lvl}"]
        except: pass
        para.paragraph_format.alignment = align_map[lvl]
        for run in para.runs:
            _force_run_formatting(run, hex_map[lvl], size_map[lvl],
                                  bold=True, italic=(lvl==3), all_caps=(lvl==1))
        pf = para.paragraph_format
        pf.space_before   = {1: Pt(24), 2: Pt(16), 3: Pt(12)}[lvl]
        pf.space_after    = {1: Pt(8),  2: Pt(6),  3: Pt(4)}[lvl]
        pf.keep_with_next = True
        if lvl == 1:
            _set_para_border(para, "bottom", HEX_H1, "8")


# ── Estimation numéros de page ────────────────────────────────────────────────
def _estimate_page_numbers(doc: Document, heading_map: dict, start_page: int = 1) -> dict:
    """
    Estime le numéro de page de chaque titre.
    Compte les caractères et détecte les sauts de page explicites.
    """
    page_map     = {}
    char_count   = 0
    current_page = start_page

    for i, para in enumerate(doc.paragraphs):
        # Détecter saut de page explicite
        has_pb = False
        for br in para._element.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                has_pb = True
                break
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            pbB = pPr.find(qn("w:pageBreakBefore"))
            if pbB is not None and pbB.get(qn("w:val"), "true") != "false":
                has_pb = True

        if has_pb:
            current_page += 1
            char_count    = 0

        # Enregistrer la page du titre
        if i in heading_map:
            page_map[i] = current_page

        # Accumuler caractères
        char_count += len(para.text) + 1

        # Saut de page naturel basé sur volume
        if char_count >= CHARS_PER_PAGE:
            current_page += char_count // CHARS_PER_PAGE
            char_count    = char_count % CHARS_PER_PAGE

        # H1 keep_with_next — souvent en haut si proche de la fin de page
        if heading_map.get(i) == 1 and char_count > CHARS_PER_PAGE * 0.85:
            current_page += 1
            char_count    = 0

    return page_map


# ── TOC manuelle avec pages estimées ─────────────────────────────────────────
def _add_toc_manual(doc: Document, heading_map: dict, start_page: int = 1) -> bool:
    if not heading_map:
        return False

    # Estimer les numéros de page AVANT d'insérer la TOC
    page_map = _estimate_page_numbers(doc, heading_map, start_page)

    toc_entries = []
    for i, para in enumerate(doc.paragraphs):
        lvl = heading_map.get(i)
        if lvl:
            toc_entries.append({
                "level": min(lvl, 3),
                "text":  para.text.strip(),
                "page":  page_map.get(i, start_page),
            })

    if not toc_entries:
        return False

    first = next((p for p in doc.paragraphs if p.text.strip()), None)
    toc_elements = []

    # Titre centré avec filet
    title_p   = OxmlElement("w:p")
    title_ppr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "480"); sp.set(qn("w:after"), "360")
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "4");    bot.set(qn("w:color"), HEX_H1)
    pBdr.append(bot)
    title_ppr.append(jc); title_ppr.append(sp); title_ppr.append(pBdr)
    title_p.append(title_ppr)
    title_r   = OxmlElement("w:r")
    title_rpr = OxmlElement("w:rPr")
    for tag, val in [("w:b", None), ("w:caps", None), ("w:sz", "32"), ("w:color", HEX_H1)]:
        el = OxmlElement(tag)
        if val: el.set(qn("w:val"), val)
        title_rpr.append(el)
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), FONT_HEAD); rf.set(qn("w:hAnsi"), FONT_HEAD)
    title_rpr.append(rf)
    title_r.append(title_rpr)
    title_t = OxmlElement("w:t"); title_t.text = "Table des matieres"
    title_r.append(title_t); title_p.append(title_r)
    toc_elements.append(title_p)

    level_cfg = {
        1: {"indent": "0",   "font_size": "22", "bold": True,  "color": HEX_H1, "before": "120", "after": "60"},
        2: {"indent": "360", "font_size": "20", "bold": False, "color": HEX_H2, "before": "60",  "after": "40"},
        3: {"indent": "720", "font_size": "18", "bold": False, "color": HEX_H3, "before": "40",  "after": "20"},
    }

    for entry in toc_entries:
        lvl  = entry["level"]
        text = entry["text"]
        page = entry["page"]
        cfg  = level_cfg[lvl]

        toc_p   = OxmlElement("w:p")
        toc_ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind"); ind.set(qn("w:left"), cfg["indent"])
        toc_ppr.append(ind)
        sp2 = OxmlElement("w:spacing")
        sp2.set(qn("w:before"), cfg["before"]); sp2.set(qn("w:after"), cfg["after"])
        toc_ppr.append(sp2)
        tabs = OxmlElement("w:tabs")
        tab  = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right"); tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "8640")
        tabs.append(tab); toc_ppr.append(tabs)
        toc_p.append(toc_ppr)

        # Texte du titre
        r_text  = OxmlElement("w:r")
        r_rpr   = OxmlElement("w:rPr")
        r_sz    = OxmlElement("w:sz");    r_sz.set(qn("w:val"), cfg["font_size"])
        r_col   = OxmlElement("w:color"); r_col.set(qn("w:val"), cfg["color"])
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT_HEAD); r_fonts.set(qn("w:hAnsi"), FONT_HEAD)
        r_rpr.append(r_fonts); r_rpr.append(r_sz); r_rpr.append(r_col)
        if cfg["bold"]: r_rpr.append(OxmlElement("w:b"))
        r_text.append(r_rpr)
        r_t = OxmlElement("w:t")
        r_t.set(qn("xml:space"), "preserve"); r_t.text = text
        r_text.append(r_t); toc_p.append(r_text)

        # Tab + numéro de page estimé
        r_pg     = OxmlElement("w:r")
        r_pg_rpr = OxmlElement("w:rPr")
        r_pg_sz  = OxmlElement("w:sz");    r_pg_sz.set(qn("w:val"), cfg["font_size"])
        r_pg_col = OxmlElement("w:color"); r_pg_col.set(qn("w:val"), HEX_H3)
        r_pg_f   = OxmlElement("w:rFonts")
        r_pg_f.set(qn("w:ascii"), FONT_HEAD); r_pg_f.set(qn("w:hAnsi"), FONT_HEAD)
        r_pg_rpr.append(r_pg_f); r_pg_rpr.append(r_pg_sz); r_pg_rpr.append(r_pg_col)
        r_pg.append(r_pg_rpr)
        r_pg.append(OxmlElement("w:tab"))
        pg_t = OxmlElement("w:t"); pg_t.text = str(page)
        r_pg.append(pg_t); toc_p.append(r_pg)
        toc_elements.append(toc_p)

    # Saut de page
    pb_p = OxmlElement("w:p")
    pb_r = OxmlElement("w:r")
    pb_e = OxmlElement("w:br"); pb_e.set(qn("w:type"), "page")
    pb_r.append(pb_e); pb_p.append(pb_r)
    toc_elements.append(pb_p)

    if first is not None:
        ref = first._element
        par = ref.getparent()
        idx = list(par).index(ref)
        for el in reversed(toc_elements):
            par.insert(idx, el)
    else:
        for el in toc_elements:
            doc.element.body.append(el)

    logger.info(f"TOC inseree — {len(toc_entries)} entrees avec pages estimees")
    return True


# ── Helpers XML ───────────────────────────────────────────────────────────────
def _set_run_color(run, hex_color):
    rpr = run._r.get_or_add_rPr()
    c = rpr.find(qn("w:color"))
    if c is None:
        c = OxmlElement("w:color"); rpr.append(c)
    c.set(qn("w:val"), hex_color)
    for attr in [qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade")]:
        if attr in c.attrib: del c.attrib[attr]


def _force_run_formatting(run, hex_color, size, bold=True, italic=False, all_caps=False):
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    for a in ["w:ascii", "w:hAnsi", "w:cs"]: rf.set(qn(a), FONT_HEAD)
    half = str(int(size.pt * 2))
    for tag in ["w:sz", "w:szCs"]:
        el = rpr.find(qn(tag))
        if el is None: el = OxmlElement(tag); rpr.append(el)
        el.set(qn("w:val"), half)
    b = rpr.find(qn("w:b"))
    if bold and b is None:           rpr.append(OxmlElement("w:b"))
    elif not bold and b is not None: rpr.remove(b)
    i_el = rpr.find(qn("w:i"))
    if italic and i_el is None:           rpr.append(OxmlElement("w:i"))
    elif not italic and i_el is not None: rpr.remove(i_el)
    caps_el = rpr.find(qn("w:caps"))
    if all_caps and caps_el is None:           rpr.append(OxmlElement("w:caps"))
    elif not all_caps and caps_el is not None: rpr.remove(caps_el)
    _set_run_color(run, hex_color)
    hl = rpr.find(qn("w:highlight"))
    if hl is not None: rpr.remove(hl)


def _set_para_border(para, side, color_hex, sz="6"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el   = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single"); el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "4");    el.set(qn("w:color"), color_hex)
    pBdr.append(el)
    old = pPr.find(qn("w:pBdr"))
    if old is not None: pPr.remove(old)
    pPr.append(pBdr)


# ── Tableaux ──────────────────────────────────────────────────────────────────
def _format_tables(doc) -> int:
    count = 0
    for table in doc.tables:
        try:
            for row_idx, row in enumerate(table.rows):
                is_header = (row_idx == 0)
                is_alt    = (row_idx % 2 == 0) and not is_header
                for cell in row.cells:
                    bg = HEX_TBL_HEAD if is_header else (HEX_TBL_ALT if is_alt else "FFFFFF")
                    _set_cell_bg(cell, bg)
                    _set_cell_borders(cell)
                    for para in cell.paragraphs:
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after  = Pt(2)
                        for run in para.runs:
                            run.font.name = FONT_BODY
                            run.font.size = Pt(10)
                            run.font.bold = is_header
                            _set_run_color(run, "FFFFFF" if is_header else HEX_BODY)
            count += 1
        except Exception as e:
            logger.warning(f"Tableau ignore: {e}")
    return count


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB  = tcPr.find(qn("w:tcBorders"))
    if tcB is None: tcB = OxmlElement("w:tcBorders"); tcPr.append(tcB)
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = tcB.find(qn(f"w:{side}"))
        if el is None: el = OxmlElement(f"w:{side}"); tcB.append(el)
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0");    el.set(qn("w:color"), "BFBFBF")


# ── Images ────────────────────────────────────────────────────────────────────
def _fix_images(doc) -> int:
    MAX = Cm(14)
    NS  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    n   = 0
    for para in doc.paragraphs:
        has_img = False
        for run in para.runs:
            for tag in [f"{{{NS}}}inline", f"{{{NS}}}anchor"]:
                for d in run._element.findall(f".//{tag}"):
                    has_img = True
                    ext = d.find(f"{{{NS}}}extent")
                    if ext is not None:
                        cx = int(ext.get("cx", 0))
                        mx = int(MAX.emu)
                        if cx > mx:
                            cy = int(ext.get("cy", 0))
                            ext.set("cx", str(mx))
                            ext.set("cy", str(int(cy * mx / cx)))
                            n += 1
        if has_img:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return n


# ── Pagination ────────────────────────────────────────────────────────────────
def _add_pagination(doc, start_page=1):
    if not doc.sections:
        return
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    section = doc.sections[-1]
    if start_page != 1:
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn("w:pgNumType"))
        if pgNumType is None:
            pgNumType = OxmlElement("w:pgNumType"); sectPr.append(pgNumType)
        pgNumType.set(qn("w:start"), str(start_page))

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_border(fp, "top", "CCCCCC", "4")
    r1 = fp.add_run("Page ")
    r1.font.name = FONT_BODY; r1.font.size = Pt(9)
    _set_run_color(r1, HEX_FOOTER)
    _add_field(fp, "PAGE")
    r2 = fp.add_run(" / ")
    r2.font.name = FONT_BODY; r2.font.size = Pt(9)
    _set_run_color(r2, HEX_FOOTER)
    _add_field(fp, "NUMPAGES")


def _add_field(para, name):
    for ftype, content in [("begin", None), (None, name), ("end", None)]:
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        sz  = OxmlElement("w:sz");    sz.set(qn("w:val"), "18")
        col = OxmlElement("w:color"); col.set(qn("w:val"), HEX_FOOTER)
        rpr.append(sz); rpr.append(col); r.append(rpr)
        if ftype:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), ftype); r.append(fc)
        else:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve"); it.text = f" {content} "; r.append(it)
        para._p.append(r)


# ── PDF ───────────────────────────────────────────────────────────────────────
def _to_pdf(docx_path, pdf_path):
    env = {**os.environ, "HOME": "/tmp", "DISPLAY": ""}
    cmd = [
        settings.LIBREOFFICE_PATH,
        "--headless", "--norestore", "--nofirststartwizard",
        "--convert-to", "pdf",
        "--outdir", str(pdf_path.parent),
        str(docx_path),
    ]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=180, env=env)
        if r.stderr: logger.warning(f"LibreOffice: {r.stderr[:200]}")
        generated = pdf_path.parent / (docx_path.stem + ".pdf")
        if generated.exists() and generated != pdf_path:
            generated.rename(pdf_path)
        if not pdf_path.exists():
            raise FileNotFoundError("PDF non genere")
        logger.info(f"PDF OK: {pdf_path.stat().st_size // 1024} Ko")
    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice timeout 180s")
    except FileNotFoundError:
        raise RuntimeError(f"LibreOffice introuvable: {settings.LIBREOFFICE_PATH}")


def _count_pages(doc) -> int:
    total = sum(max(1, len(p.text) // 80) for p in doc.paragraphs if p.text.strip())
    return max(1, total // 40)
